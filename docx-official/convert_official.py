#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

"""公文格式转换工具 v4（无红头，GB/T 9704标准）"""

import sys
import os
import json
import re
import requests
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

API_KEY = os.environ.get("DEEPSEEK_API_KEY")

# GB/T 9704 标准常量
FONT_TITLE = '方正小标宋简体'    # 标题字体
FONT_H1 = '黑体'                # 一级标题
FONT_H2 = '楷体_GB2312'         # 二级标题
FONT_BODY = '仿宋_GB2312'       # 正文/三级标题/四级标题
FONT_TABLE_HEADER = '黑体'      # 表头
FONT_TABLE_BODY = '宋体'        # 表格正文

SIZE_TITLE = 22    # 二号
SIZE_BODY = 16     # 三号
SIZE_TABLE_H = 10.5  # 五号
SIZE_TABLE_B = 10.5  # 五号
SIZE_PAGE_NUM = 14   # 四号

LINE_SPACING = 28   # 公文固定行距28磅


def clean_text(text):
    """清洗文本：去除多余空格、处理从网页复制的特殊字符"""
    if not text:
        return text

    # 替换特殊空白字符
    text = text.replace('\u3000', ' ')   # 全角空格
    text = text.replace('\xa0', ' ')     # 不间断空格
    text = text.replace('\u200b', '')    # 零宽空格
    text = text.replace('\u200c', '')    # 零宽非连接符
    text = text.replace('\u200d', '')    # 零宽连接符
    text = text.replace('\ufeff', '')    # BOM
    text = text.replace('\t', ' ')       # 制表符

    # 去除多余空格（保留单个空格）
    text = re.sub(r' {2,}', ' ', text)

    # 去除中文字符之间的空格
    text = re.sub(r'([\u4e00-\u9fa5，。、；：""''！？（）])\s+([\u4e00-\u9fa5，。、；：""''！？（）])', r'\1\2', text)
    text = re.sub(r'([\u4e00-\u9fa5，。、；：""''！？（）])\s+([\u4e00-\u9fa5，。、；：""''！？（）])', r'\1\2', text)

    # 去除行首行尾空格
    text = text.strip()

    return text


def normalize_paragraphs(doc):
    """规范化段落：清洗文本，去除空行"""
    cleaned = []
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if text:
            cleaned.append(text)
    return cleaned


def set_run_font(run, font_name):
    """设置run的中文字体（eastAsia）和西文字体"""
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_first_line_chars(paragraph, chars=200):
    """设置首行缩进（按字符数，GB/T 9704标准用200=2字符）"""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars))


def check_fonts():
    """检查系统字体（通过字体文件是否存在判断）"""
    fonts_dir = os.path.join(os.environ.get('WINDIR', r'C:\Windows'), 'Fonts')
    font_files = {
        '方正小标宋简体': 'FZSTK.TTF',
        '楷体': 'simkai.ttf',
        '仿宋': 'simfang.ttf',
        '黑体': 'simhei.ttf'
    }
    result = {}
    for font, filename in font_files.items():
        result[font] = os.path.exists(os.path.join(fonts_dir, filename))
    return result


def set_page_margins(section):
    """设置页边距（GB/T 9704标准）"""
    section.page_width = Cm(21.0)    # A4
    section.page_height = Cm(29.7)   # A4
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def correct_text_with_api(texts, api_key=None):
    """使用DeepSeek API检查并修正文本中的错别字、漏字、明显错误"""
    if not api_key:
        api_key = API_KEY

    if not api_key:
        return texts, []

    # 准备文本列表
    text_list = []
    for i, text in enumerate(texts):
        if text and len(text) > 2:
            text_list.append(f"[{i}] {text}")

    if not text_list:
        return texts, []

    content = "\n".join(text_list)

    prompt = f"""请检查以下公文段落中的错误，包括：
1. 错别字（同音字、形近字错误）
2. 漏字（缺少必要的字，如"日"字缺失）
3. 明显的事实或格式错误（如日期格式不完整）
4. 标点符号错误

注意：
- 只修改确实有错误的地方，不要改变原文的表述风格
- 不要修改专业术语或特定表述
- 不要添加原文没有的内容

段落列表：
{content}

请输出JSON格式，只列出需要修改的段落：
{{
  "corrections": [
    {{"index": 段落索引, "original": "原文片段", "corrected": "修正后片段", "reason": "修改原因"}}
  ]
}}

如果没有需要修改的内容，返回：{{"corrections": []}}
只输出JSON，不要其他内容。"""

    try:
        print("调用DeepSeek API检查文本错误...")
        response = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": "你是公文校对专家。只输出JSON，不要输出任何解释。只修改确实有错误的地方。"},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.1,
                "max_tokens": 2000
            },
            timeout=60
        )

        result = response.json()

        if "error" in result:
            print(f"API错误: {result['error']['message']}")
            return texts, []

        raw_text = result["choices"][0]["message"]["content"].strip()

        # 提取JSON
        if "```json" in raw_text:
            raw_text = raw_text.split("```json")[1].split("```")[0]
        elif "```" in raw_text:
            raw_text = raw_text.split("```")[1].split("```")[0]

        data = json.loads(raw_text)
        corrections = data.get('corrections', [])

        if not corrections:
            print("未发现明显错误")
            return texts, []

        # 应用修正
        corrected_texts = texts.copy()
        changes = []
        for c in corrections:
            idx = c.get('index')
            original = c.get('original', '')
            corrected = c.get('corrected', '')
            reason = c.get('reason', '')

            if idx is not None and idx < len(corrected_texts) and original:
                if original in corrected_texts[idx]:
                    corrected_texts[idx] = corrected_texts[idx].replace(original, corrected, 1)
                    changes.append({
                        'index': idx,
                        'original': original,
                        'corrected': corrected,
                        'reason': reason
                    })

        print(f"发现并修正了 {len(changes)} 处错误")
        return corrected_texts, changes

    except Exception as e:
        print(f"文本纠错失败: {e}")
        return texts, []


def analyze_with_api(texts, api_key=None):
    """使用DeepSeek API分析文档结构（只返回类型分类，不返回文本）"""
    if not api_key:
        api_key = API_KEY

    if not api_key:
        print("提示：未设置DEEPSEEK_API_KEY，使用规则判断结构")
        return None

    # 准备文本摘要（只用于分析，不用于最终输出）
    text_list = []
    for i, text in enumerate(texts):
        if text and len(text) > 1:
            text_list.append({"index": i, "text": text[:150]})

    if not text_list:
        return None

    content = json.dumps(text_list, ensure_ascii=False, indent=2)

    prompt = f"""分析以下文档段落，为每个段落标注类型。

类型说明：
- "title": 文档主标题（通常是第一个较短的居中标题行）
- "h1": 一级标题（以"一、""二、""三、"等中文数字+顿号开头）
- "h2": 二级标题（以"（一）""（二）"等中文数字+括号开头）
- "h3": 三级标题（以"1.""2.""3."等阿拉伯数字+点开头）
- "h4": 四级标题（以"（1）""（2）"等阿拉伯数字+括号开头）
- "date": 落款日期行（含"执行日期""印发日期""成文日期"等，或文末的"XXXX年X月X日"格式）
- "signature": 落款单位名称行（如"XX学校""XX单位"，通常在文档最后几行，可能与日期同行）
- "body": 正文内容

注意：
1. 只输出JSON，不要输出任何其他内容。不要修改或重写段落原文。
2. 文档最后2-3行通常是落款，需仔细识别。
3. 如果一行同时包含单位名称和日期（如"钢城现代学校 2026年5月12"），标为"signature"。

段落列表：
{content}

输出格式（严格JSON）：
{{
  "title": "主标题原文（从原文复制，不要修改）",
  "types": [
    {{"index": 0, "type": "title"}},
    {{"index": 1, "type": "body"}},
    ...
  ]
}}"""

    try:
        print("调用DeepSeek API分析文档结构...")
        response = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": "你是文档结构分析助手。只输出JSON分类结果，不要输出任何解释或额外文本。不要修改或重写段落原文。"},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.1,
                "max_tokens": 2000
            },
            timeout=60
        )

        result = response.json()

        if "error" in result:
            print(f"API错误: {result['error']['message']}")
            return None

        raw_text = result["choices"][0]["message"]["content"].strip()

        # 提取JSON
        if "```json" in raw_text:
            raw_text = raw_text.split("```json")[1].split("```")[0]
        elif "```" in raw_text:
            raw_text = raw_text.split("```")[1].split("```")[0]

        data = json.loads(raw_text)

        # 构建sections：使用API的type分类，但使用原始cleaned_texts的文本
        types_map = {}
        for item in data.get('types', []):
            types_map[item['index']] = item['type']

        sections = []
        for i, text in enumerate(texts):
            ptype = types_map.get(i, 'body')
            sections.append({"index": i, "type": ptype, "text": text})

        print(f"API分析完成：{len(sections)} 个段落")
        return {"title": data.get('title', ''), "sections": sections}

    except Exception as e:
        print(f"API调用失败: {e}")
        return None


def analyze_by_rules(texts):
    """使用规则分析文档结构"""
    sections = []
    total = len(texts)

    for i, text in enumerate(texts):
        ptype = 'body'

        # 一级标题：一、二、三、...
        if re.match(r'^[一二三四五六七八九十]+、', text) and len(text) < 50:
            ptype = 'h1'
        # 二级标题：（一）（二）（三）...
        elif re.match(r'^（[一二三四五六七八九十]+）', text) and len(text) < 50:
            ptype = 'h2'
        # 三级标题：1. 2. 3. ...
        elif re.match(r'^\d+[.、]', text) and len(text) < 50:
            ptype = 'h3'
        # 四级标题：（1）（2）（3）...
        elif re.match(r'^（\d+）', text) and len(text) < 50:
            ptype = 'h4'
        # 落款日期行（含"执行日期""印发日期"等）
        elif re.search(r'(执行日期|印发日期|成文日期|发布日期)', text):
            ptype = 'date'
            # 转换中文日期为阿拉伯数字
            text = convert_chinese_date_to_arabic(text)
        # 纯日期行（阿拉伯数字年月日格式，且在文末5段内）
        elif i >= total - 5 and re.search(r'\d{4}年\d{1,2}月\d{1,2}日', text):
            ptype = 'date'
        # 中文日期行（文末5段内）
        elif i >= total - 5 and re.search(r'[一二〇零三四五六七八九十]+年[一二三四五六七八九十]+月', text):
            ptype = 'date'
            text = convert_chinese_date_to_arabic(text)
        # 落款单位+日期（文末较短行，含年月）
        elif i >= total - 3 and len(text) < 40 and re.search(r'\d{4}年', text):
            ptype = 'signature'
        # 落款单位+中文日期（文末较短行）
        elif i >= total - 3 and len(text) < 40 and re.search(r'[一二〇零三四五六七八九十]+年', text):
            ptype = 'signature'
            text = convert_chinese_date_to_arabic(text)

        sections.append({"index": i, "type": ptype, "text": text})

    return {"sections": sections}


def add_paragraph_with_format(new_doc, text, ptype, font_title, font_body, font_header, font_unit):
    """添加一个段落并设置公文格式"""
    p = new_doc.add_paragraph()

    # 设置字体和字号
    run = p.add_run(text)
    run.font.size = Pt(SIZE_BODY)

    if ptype == 'title':
        set_run_font(run, font_title)
        run.font.size = Pt(SIZE_TITLE)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(30)
    elif ptype == 'h1':
        set_run_font(run, font_header)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_SPACING)
        set_first_line_chars(p, 200)
    elif ptype == 'h2':
        set_run_font(run, font_unit)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_SPACING)
        set_first_line_chars(p, 200)
    elif ptype == 'h3':
        set_run_font(run, font_body)
        run.bold = True
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_SPACING)
        set_first_line_chars(p, 200)
    elif ptype == 'h4':
        set_run_font(run, font_body)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_SPACING)
        set_first_line_chars(p, 200)
    elif ptype == 'date':
        set_run_font(run, font_body)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_indent = Pt(64)  # 右空四字（4×16pt）
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_SPACING)
    elif ptype == 'signature':
        set_run_font(run, font_body)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_indent = Pt(64)  # 右空四字（4×16pt）
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_SPACING)
    else:  # body
        set_run_font(run, font_body)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_SPACING)
        set_first_line_chars(p, 200)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    return p


def add_page_numbers(doc):
    """添加页码（宋体小四号，页脚外侧，-X-格式，左右空1字符）"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中

        # 左空1字符
        run_space1 = p.add_run(' ')
        run_space1.font.name = '宋体'
        run_space1.font.size = Pt(12)  # 小四号
        r_s1 = run_space1._element
        rPr_s1 = r_s1.find(qn('w:rPr'))
        if rPr_s1 is None:
            rPr_s1 = OxmlElement('w:rPr')
            r_s1.insert(0, rPr_s1)
        rFonts_s1 = rPr_s1.find(qn('w:rFonts'))
        if rFonts_s1 is None:
            rFonts_s1 = OxmlElement('w:rFonts')
            rPr_s1.insert(0, rFonts_s1)
        rFonts_s1.set(qn('w:eastAsia'), '宋体')

        # 前缀 "-"
        run1 = p.add_run('-')
        run1.font.name = '宋体'
        run1.font.size = Pt(12)  # 小四号
        r1 = run1._element
        rPr1 = r1.find(qn('w:rPr'))
        if rPr1 is None:
            rPr1 = OxmlElement('w:rPr')
            r1.insert(0, rPr1)
        rFonts1 = rPr1.find(qn('w:rFonts'))
        if rFonts1 is None:
            rFonts1 = OxmlElement('w:rFonts')
            rPr1.insert(0, rFonts1)
        rFonts1.set(qn('w:eastAsia'), '宋体')

        # 页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_field = p.add_run()
        run_field._element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run_instr = p.add_run()
        run_instr._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_end = p.add_run()
        run_end._element.append(fldChar2)

        # 后缀 "-"
        run2 = p.add_run('-')
        run2.font.name = '宋体'
        run2.font.size = Pt(12)  # 小四号
        r2 = run2._element
        rPr2 = r2.find(qn('w:rPr'))
        if rPr2 is None:
            rPr2 = OxmlElement('w:rPr')
            r2.insert(0, rPr2)
        rFonts2 = rPr2.find(qn('w:rFonts'))
        if rFonts2 is None:
            rFonts2 = OxmlElement('w:rFonts')
            rPr2.insert(0, rFonts2)
        rFonts2.set(qn('w:eastAsia'), '宋体')

        # 右空1字符
        run_space2 = p.add_run(' ')
        run_space2.font.name = '宋体'
        run_space2.font.size = Pt(12)  # 小四号
        r_s2 = run_space2._element
        rPr_s2 = r_s2.find(qn('w:rPr'))
        if rPr_s2 is None:
            rPr_s2 = OxmlElement('w:rPr')
            r_s2.insert(0, rPr_s2)
        rFonts_s2 = rPr_s2.find(qn('w:rFonts'))
        if rFonts_s2 is None:
            rFonts_s2 = OxmlElement('w:rFonts')
            rPr_s2.insert(0, rFonts_s2)
        rFonts_s2.set(qn('w:eastAsia'), '宋体')


def split_signature_date(text):
    """拆分落款行：如果同时包含单位名称和日期，拆成两行，并转换为阿拉伯数字日期"""
    # 匹配模式：单位名称 + 空格 + 日期
    match = re.match(r'^(.+?)\s+(\d{4}年\d{1,2}月\d{1,2}日?)$', text)
    if match:
        return match.group(1), convert_date_to_arabic(match.group(2))

    # 匹配模式：单位名称 + 空格 + 日期（无"日"）
    match = re.match(r'^(.+?)\s+(\d{4}年\d{1,2}月\d{1,2})$', text)
    if match:
        return match.group(1), convert_date_to_arabic(match.group(2))

    # 匹配模式：单位名称 + 空格 + 中文日期
    match = re.match(r'^(.+?)\s+([一二〇零三四五六七八九十]+年[一二三四五六七八九十]+月[一二三四五六七八九十]+日?)$', text)
    if match:
        return match.group(1), convert_chinese_date_to_arabic(match.group(2))

    return None, None


def convert_date_to_arabic(date_str):
    """确保日期为阿拉伯数字格式"""
    # 已经是阿拉伯数字格式，直接返回
    if re.match(r'^\d{4}年\d{1,2}月\d{1,2}日?$', date_str):
        return date_str
    return date_str


def convert_chinese_date_to_arabic(date_str):
    """将中文数字日期转换为阿拉伯数字日期"""
    cn_num = {
        '〇': '0', '零': '0', '一': '1', '二': '2', '三': '3', '四': '4',
        '五': '5', '六': '6', '七': '7', '八': '8', '九': '9', '十': '10',
        '十一': '11', '十二': '12'
    }

    result = date_str

    # 转换年份
    year_match = re.search(r'([一二〇零三四五六七八九十]+)年', result)
    if year_match:
        cn_year = year_match.group(1)
        ar_year = ''
        for ch in cn_year:
            ar_year += cn_num.get(ch, ch)
        result = result.replace(cn_year + '年', ar_year + '年')

    # 转换月份
    month_match = re.search(r'([一二三四五六七八九十]+)月', result)
    if month_match:
        cn_month = month_match.group(1)
        ar_month = cn_num.get(cn_month, cn_month)
        result = result.replace(cn_month + '月', ar_month + '月')

    # 转换日期
    day_match = re.search(r'([一二三四五六七八九十]+)日', result)
    if day_match:
        cn_day = day_match.group(1)
        ar_day = cn_num.get(cn_day, cn_day)
        result = result.replace(cn_day + '日', ar_day + '日')

    return result


def convert_to_official(input_file, output_file=None, use_api=True):
    """转换为公文格式（无红头，GB/T 9704标准）"""
    if not output_file:
        stem = Path(input_file).stem
        output_file = str(Path(input_file).parent / f"{stem}_公文格式.docx")

    print(f"\n{'='*50}")
    print(f"读取: {input_file}")

    # 检查字体
    print("\n检查系统字体...")
    fonts_status = check_fonts()
    missing = [f for f, ok in fonts_status.items() if not ok]
    if missing:
        print(f"⚠️ 缺少字体: {', '.join(missing)}，将使用替代字体")

    # 读取原文档
    doc = Document(input_file)

    # 清洗段落
    print("清洗文本...")
    cleaned_texts = normalize_paragraphs(doc)
    print(f"原始段落数: {len(doc.paragraphs)}, 清洗后: {len(cleaned_texts)}")

    # 文本纠错（使用API检查错别字、漏字等）
    changes = []
    if use_api and API_KEY:
        cleaned_texts, changes = correct_text_with_api(cleaned_texts)

    # 分析结构
    if use_api and API_KEY:
        structure = analyze_with_api(cleaned_texts)
    else:
        structure = None

    if not structure:
        print("使用规则判断文档结构...")
        structure = analyze_by_rules(cleaned_texts)

    # 获取可用字体
    font_title = FONT_TITLE if fonts_status.get('方正小标宋简体') else 'SimSun'
    font_body = FONT_BODY if fonts_status.get('仿宋') else 'FangSong'
    font_header = FONT_H1 if fonts_status.get('黑体') else 'SimHei'
    font_unit = FONT_H2 if fonts_status.get('楷体') else 'KaiTi'

    print(f"\n使用字体：")
    print(f"  标题: {font_title}")
    print(f"  正文: {font_body}")
    print(f"  一级标题: {font_header}")
    print(f"  二级标题: {font_unit}")

    # 创建新文档
    new_doc = Document()
    new_section = new_doc.sections[0]
    set_page_margins(new_section)

    # 处理标题
    title_text = structure.get('title', '')
    sections = structure.get('sections', [])

    # 如果没有从API获取到标题，尝试从前3段找
    if not title_text:
        for item in sections[:3]:
            t = item.get('text', '')
            if 4 < len(t) < 30 and item.get('type') == 'body':
                title_text = t
                item['type'] = 'title'
                break

    # 写入正文
    for item in sections:
        text = item.get('text', '')
        ptype = item.get('type', 'body')

        if not text:
            continue

        # 标题单独处理
        if ptype == 'title' and text == title_text:
            add_paragraph_with_format(new_doc, text, 'title', font_title, font_body, font_header, font_unit)
            continue

        # 日期段落：转换中文日期为阿拉伯数字
        if ptype == 'date' and re.search(r'[一二〇零三四五六七八九十]+年', text):
            text = convert_chinese_date_to_arabic(text)

        # 落款处理：如果同时包含单位和日期，拆成两行
        if ptype == 'signature':
            unit, date = split_signature_date(text)
            if unit and date:
                add_paragraph_with_format(new_doc, unit, 'signature', font_title, font_body, font_header, font_unit)
                add_paragraph_with_format(new_doc, date, 'date', font_title, font_body, font_header, font_unit)
                continue

        add_paragraph_with_format(new_doc, text, ptype, font_title, font_body, font_header, font_unit)

    # 添加页码
    add_page_numbers(new_doc)

    # 表格处理
    if doc.tables and len(doc.tables) > 0:
        print(f"\n识别到 {len(doc.tables)} 个表格")
        for idx, table in enumerate(doc.tables):
            # 表格标题
            p_t = new_doc.add_paragraph()
            p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_t = p_t.add_run(f"表{idx + 1}")
            set_run_font(r_t, font_header)
            r_t.font.size = Pt(SIZE_TABLE_H)
            r_t.bold = True
            p_t.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p_t.paragraph_format.line_spacing = Pt(LINE_SPACING)

            # 创建表格
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            new_table.style = 'Table Grid'

            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell_text = clean_text(cell.text)
                    target = new_table.rows[i].cells[j]
                    target.text = cell_text
                    for para in target.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        para.paragraph_format.line_spacing = Pt(14)
                        for run in para.runs:
                            run.font.size = Pt(SIZE_TABLE_B)
                            if i == 0:
                                set_run_font(run, FONT_TABLE_HEADER)
                                run.bold = True
                            else:
                                set_run_font(run, FONT_TABLE_BODY)

    # 保存
    new_doc.save(output_file)
    print(f"\n✅ 已保存: {output_file}")

    # 报告文本修正
    if changes:
        print(f"\n{'='*50}")
        print("📝 文本纠错报告：")
        for c in changes:
            print(f"  段落{c['index']+1}: \"{c['original']}\" → \"{c['corrected']}\"")
            print(f"    原因: {c['reason']}")
        print(f"{'='*50}")

    print(f"\n提示：页码已自动添加（-X-格式，居中）")
    print(f"{'='*50}\n")

    return output_file


def main():
    if len(sys.argv) < 2:
        print("用法: python convert_official.py <输入文件.docx> [输出文件.docx] [--no-api]")
        print("\n参数:")
        print("  --no-api  不使用API（纯规则判断）")
        sys.exit(1)

    args = [arg for arg in sys.argv[1:] if not arg.startswith('--')]
    use_api = "--no-api" not in sys.argv

    input_file = args[0]
    output_file = args[1] if len(args) > 1 else None

    if not os.path.exists(input_file):
        print(f"错误：文件不存在: {input_file}")
        sys.exit(1)

    try:
        convert_to_official(input_file, output_file, use_api)
    except Exception as e:
        print(f"错误：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
